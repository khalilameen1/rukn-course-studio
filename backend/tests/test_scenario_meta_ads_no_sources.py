"""End-to-end test scenario: a real Arabic course brief, no uploaded
sources, connected modules with bridge projects, no manual map.

Runs with FakeProvider (always) and, only if a real ANTHROPIC_API_KEY is
configured, also with AnthropicProvider - skipped otherwise rather than
failing, since a key is an environment concern, not a code concern.

Checks the specific guarantees this scenario was written to prove:
- the pipeline never raises out of run_generation (no crash)
- a real .docx is produced and downloadable
- the job's API-facing shape never carries reel content or internal logs
- job status/stage/progress reporting works throughout
- the downloaded DOCX itself follows the teleprompter contract: course
  title, module heading, lesson heading, spoken script - and none of the
  internal review/validation/quality-check machinery this same pipeline
  run generates internally along the way (see docs/PRD.md and
  rukn_teleprompter_docx_contract in app/seed_admin_knowledge.py)
"""

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlmodel import Session, SQLModel, create_engine

from app.ai.anthropic_provider import AnthropicProvider
from app.ai.fake_provider import FakeProvider
from app.config import settings
from app.crud import course_versions, courses
from app.generation.orchestrator import run_generation
from app.models.enums import ExplanationLevel, JobStatus, StructureMode
from app.schemas.generation_job import GenerationJobRead

# Mirrors the FORBIDDEN_SUBSTRINGS list in test_docx_export.py - kept
# duplicated (not imported) so this end-to-end test still fails loudly on
# its own if that file's list is ever narrowed.
FORBIDDEN_DOCX_SUBSTRINGS = [
    "internal_review",
    "validation",
    "quality_check",
    "prepared by ai",
    "methodology",
    "note to instructor",
    "say this",
    "explain that",
]

COURSE_BRIEF = dict(
    title="تصميم إعلانات ميتا بالفوتوشوب",
    audience="مبتدئ يريد تصميم إعلانات لمطاعم وعيادات وصفحات صغيرة",
    outcome="يخرج قادرًا على تصميم إعلان مفهوم وقابل للاختبار",
    structure_mode=StructureMode.CONNECTED_MODULES_WITH_BRIDGE_PROJECTS,
    manual_map_text=None,  # explicitly empty - no manual map, no sources
    explanation_level=ExplanationLevel.FINAL_ONLY,
)


@pytest.fixture()
def session(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as s:
        yield s


@pytest.fixture(autouse=True)
def isolated_storage(tmp_path, monkeypatch):
    import app.generation.orchestrator as orchestrator_module

    monkeypatch.setattr(orchestrator_module.settings, "storage_outputs_dir", tmp_path)


def _make_course(session):
    return courses.create(session, **COURSE_BRIEF)


def _assert_job_never_leaks_internal_content(job):
    """Simulates what the API actually returns (GenerationJobRead) and
    checks reel/log content never appears in it - this is the concrete,
    checkable form of "internal reels are not shown in frontend"."""
    api_shape = GenerationJobRead.model_validate(job).model_dump()

    assert "log_json" not in api_shape
    serialized = str(api_shape)
    assert "script_text" not in serialized
    assert "used_ideas" not in serialized
    assert "used_examples" not in serialized

    # Sanity: log_json really does exist on the underlying job (proving the
    # schema is excluding it deliberately, not because it's simply absent).
    assert isinstance(job.log_json, list)
    assert len(job.log_json) > 0


def test_full_generation_with_fake_provider_no_sources(session):
    course = _make_course(session)

    job = run_generation(session, course.id, provider=FakeProvider())

    # No crash: run_generation always returns a job, never raises for a
    # provider-level failure (see docs/PRD.md FR-10).
    assert job is not None
    assert job.status == JobStatus.COMPLETED
    assert job.current_stage == "done"
    assert job.progress_percent == 100
    assert job.error_message is None

    # A DOCX was actually created and is a real, openable file.
    assert job.output_docx_path is not None
    docx_path = job.output_docx_path
    document = Document(docx_path)
    paragraph_texts = [p.text for p in document.paragraphs]
    assert course.title in paragraph_texts

    # Teleprompter contract, checked against the real file this pipeline
    # run actually produced (not a hand-built fixture): course title,
    # module heading, lesson heading, and spoken script must all be
    # visually present.
    assert any(t.startswith("Module 1 —") for t in paragraph_texts)
    assert any(t.startswith("Lesson 1 —") for t in paragraph_texts)
    # FakeProvider's placeholder script text (app/ai/fake_provider.py) -
    # proves actual spoken-script paragraphs exist under a lesson heading,
    # not just the headings themselves. Opener wording varies with
    # lesson_curve.hook_strength; the must_cover body line is stable.
    assert any("النقطة دي مهمة" in t for t in paragraph_texts)

    # Same real file must never surface any of the internal review/
    # validation/quality-check machinery this pipeline run generated along
    # the way (see job.log_json below, which does contain review steps).
    full_text = "\n".join(paragraph_texts).lower()
    for forbidden in FORBIDDEN_DOCX_SUBSTRINGS:
        assert forbidden not in full_text

    # A CourseVersion points at that same file.
    versions = course_versions.list(session, course_id=course.id)
    assert len(versions) == 1
    assert versions[0].output_docx_path == docx_path

    # Job progress/status reporting works and never leaks reel content.
    _assert_job_never_leaks_internal_content(job)

    # Sanity: this run really did generate internal review steps - proving
    # the DOCX assertions above are actually hiding something, not just
    # trivially passing because nothing internal happened.
    review_steps = [e for e in job.log_json if "review" in e.get("step", "")]
    assert len(review_steps) > 0


def test_download_latest_docx_via_real_api_endpoints(tmp_path, monkeypatch):
    """Runs the same scenario through the actual FastAPI app (not just the
    orchestrator function directly), proving the HTTP-facing behavior:
    create course -> generate -> job status -> download - all work, and
    the job response never contains reel/log content."""
    import app.db as db_module
    import app.generation.orchestrator as orchestrator_module

    engine = create_engine(f"sqlite:///{tmp_path / 'api_test.db'}")
    SQLModel.metadata.create_all(engine)
    monkeypatch.setattr(db_module, "engine", engine)
    monkeypatch.setattr(orchestrator_module, "engine", engine)
    monkeypatch.setattr(orchestrator_module.settings, "storage_outputs_dir", tmp_path)

    from app.main import app

    client = TestClient(app)

    create_response = client.post(
        "/courses",
        json={
            "title": COURSE_BRIEF["title"],
            "audience": COURSE_BRIEF["audience"],
            "outcome": COURSE_BRIEF["outcome"],
            "structure_mode": COURSE_BRIEF["structure_mode"].value,
            "manual_map_text": None,
            "explanation_level": COURSE_BRIEF["explanation_level"].value,
        },
    )
    assert create_response.status_code == 201
    course_id = create_response.json()["id"]

    generate_response = client.post(f"/courses/{course_id}/generate")
    assert generate_response.status_code == 201
    job_body = generate_response.json()

    # Exactly the fields the frontend GeneratePanel relies on - nothing else
    # reel/log related (course_map_json/completed_reels_json/log_json stay
    # internal-only, see app/schemas/generation_job.py). `run_snapshot_json`/
    # `output_score_json`/`budget_warning` were added by the AI-ops
    # hardening pass (see app/generation/run_snapshot.py,
    # app/generation/output_scoring.py, app/generation/budget_guard.py) -
    # all three are safe-by-construction (hashes/scores/a warning string,
    # never secrets or raw source/admin-knowledge text), so they're
    # intentionally included here too.
    assert set(job_body.keys()) == {
        "id",
        "course_id",
        "status",
        "current_stage",
        "progress_percent",
        "output_docx_path",
        "error_message",
        "last_completed_step",
        "completed_modules_count",
        "completed_reels_count",
        "total_lessons_count",
        "needs_review_count",
        "error_category",
        "partial_docx_path",
        "current_module_index",
        "current_lesson_index",
        "last_progress_message",
        "last_saved_at",
        "estimated_usage_summary",
        "estimated_duration_summary",
        "internal_risk_count",
        "generation_quality_mode",
        "web_research_mode",
        "run_snapshot_json",
        "output_score_json",
        "budget_warning",
        "created_at",
        "updated_at",
        "run_status",
        "completed_lessons_count",
        "partial_docx_available",
    }
    assert job_body["status"] == "completed"
    assert job_body["current_stage"] == "done"
    assert job_body["progress_percent"] == 100
    job_id = job_body["id"]

    # Job status endpoint (what the frontend polls) works the same way.
    status_response = client.get(f"/jobs/{job_id}")
    assert status_response.status_code == 200
    assert status_response.json()["status"] == "completed"

    versions_response = client.get(f"/courses/{course_id}/versions")
    assert versions_response.status_code == 200
    assert len(versions_response.json()) == 1

    download_response = client.get(f"/courses/{course_id}/download/latest")
    assert download_response.status_code == 200
    assert download_response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(download_response.content) > 0


@pytest.mark.skipif(
    not settings.anthropic_api_key,
    reason="No ANTHROPIC_API_KEY configured - skipping real-provider run",
)
def test_full_generation_with_real_anthropic_provider_no_sources(session):
    """Only runs if a real API key is configured (per the scenario's own
    "if API key exists" condition) - this genuinely calls the Anthropic
    API when it does."""
    course = _make_course(session)

    job = run_generation(session, course.id, provider=AnthropicProvider())

    assert job is not None
    assert job.status in (JobStatus.COMPLETED, JobStatus.FAILED)
    if job.status == JobStatus.FAILED:
        pytest.fail(f"Real-provider run failed: {job.error_message}")

    assert job.output_docx_path is not None
    Document(job.output_docx_path)  # must be a real, openable docx
    _assert_job_never_leaks_internal_content(job)

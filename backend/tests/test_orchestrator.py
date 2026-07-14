"""Tests for app/generation/orchestrator.py, using FakeProvider throughout."""

import json
from dataclasses import asdict
from pathlib import Path

import pytest
from sqlmodel import Session, SQLModel, create_engine

from docx import Document

from app.ai.fake_provider import FakeProvider
from app.ai.provider import BuildCourseMapInput, WriteSingleReelInput
from app.crud import (
    admin_knowledge_items,
    course_sources,
    course_versions,
    courses,
    generation_jobs,
    source_analyses,
)
from app.generation.errors import ERROR_CATEGORY_MESSAGES, ERROR_CATEGORY_MESSAGES_NOTHING_SAVED
from app.generation.orchestrator import run_generation
from app.models.enums import (
    ExplanationLevel,
    ItemType,
    JobStatus,
    Priority,
    SourceCategory,
    StructureMode,
)
from app.services.source_analysis import SHORT_SOURCE_MAX_CHARS, analyze_source_text


@pytest.fixture()
def session(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as s:
        yield s


@pytest.fixture(autouse=True)
def isolated_storage(tmp_path, monkeypatch):
    """Redirect saved internal JSON to a temp dir instead of real storage/."""
    import app.generation.orchestrator as orchestrator_module

    monkeypatch.setattr(orchestrator_module.settings, "storage_outputs_dir", tmp_path)


def _make_course(session, **overrides):
    fields = dict(
        title="Intro to Excel Formulas",
        audience="new hires",
        outcome="build a basic budget sheet",
        structure_mode=StructureMode.CONNECTED_NO_MODULES,
        explanation_level=ExplanationLevel.FINAL_ONLY,
    )
    fields.update(overrides)
    return courses.create(session, **fields)


def _make_source_with_analysis(
    session,
    course_id,
    text,
    category=SourceCategory.SCIENTIFIC_REFERENCE,
    priority=Priority.MEDIUM,
):
    """Mirrors what app/routers/sources.py does on upload/notes creation."""
    source = course_sources.create(
        session,
        course_id=course_id,
        source_category=category,
        priority=priority,
        status="ready",
        extracted_text=text,
    )
    analysis = analyze_source_text(text, category.value)
    source_analyses.create(
        session,
        source_id=source.id,
        chunks_json=[asdict(chunk) for chunk in analysis.chunks],
        source_summary=analysis.source_summary,
        key_points_json=analysis.key_points,
        avoid_points_json=analysis.avoid_points,
    )
    return source


class RecordingProvider(FakeProvider):
    """FakeProvider that records every input it receives, so tests can
    inspect exactly what context the orchestrator sent it."""

    def __init__(self):
        super().__init__()
        self.map_calls: list[BuildCourseMapInput] = []
        self.reel_calls: list[WriteSingleReelInput] = []
        self.review_single_reel_calls = 0

    def build_course_map(self, input):  # noqa: A002 - matches AIProvider's signature
        self.map_calls.append(input)
        return super().build_course_map(input)

    def write_single_reel(self, input):  # noqa: A002
        self.reel_calls.append(input)
        return super().write_single_reel(input)

    def review_single_reel(self, input):  # noqa: A002
        self.review_single_reel_calls += 1
        return super().review_single_reel(input)


class FailAfterNReelsProvider(FakeProvider):
    """Raises once `fail_after` distinct reels have already completed
    successfully - a deterministic mid-pipeline failure, robust to any
    local-validator retries within those already-completed reels (see
    `_write_and_review_reel`'s bounded rewrite loop, which can call
    `write_single_reel` more than once for the very same reel_id)."""

    def __init__(self, fail_after: int, message: str = "simulated provider failure"):
        super().__init__()
        self.fail_after = fail_after
        self.message = message
        self._seen_reel_ids: set[str] = set()

    def write_single_reel(self, input):  # noqa: A002
        if input.reel.reel_id not in self._seen_reel_ids and len(self._seen_reel_ids) >= self.fail_after:
            raise RuntimeError(self.message)
        self._seen_reel_ids.add(input.reel.reel_id)
        return super().write_single_reel(input)


class RealisticRateLimitError(Exception):
    """Named and worded like the real `anthropic` SDK's `RateLimitError`
    (e.g. `anthropic.RateLimitError: Error code: 429 - rate limit
    exceeded`) without importing the actual `anthropic` package here -
    `classify_provider_error` (app/generation/errors.py) is a pure
    function that only looks at exception type name/message text, so this
    is enough to prove it classifies a real-shaped failure correctly."""


class RealisticRateLimitProvider(FakeProvider):
    """Mirrors `FailAfterNReelsProvider` below, but raises an
    Anthropic-shaped exception instead of a generic `RuntimeError`."""

    def __init__(self, fail_after: int):
        super().__init__()
        self.fail_after = fail_after
        self._seen_reel_ids: set[str] = set()

    def write_single_reel(self, input):  # noqa: A002
        if input.reel.reel_id not in self._seen_reel_ids and len(self._seen_reel_ids) >= self.fail_after:
            raise RealisticRateLimitError("Error code: 429 - rate limit exceeded")
        self._seen_reel_ids.add(input.reel.reel_id)
        return super().write_single_reel(input)


def test_run_generation_completes_end_to_end(session):
    course = _make_course(session)

    job = run_generation(session, course.id)

    assert job.status == JobStatus.COMPLETED
    assert job.current_stage == "done"
    assert job.progress_percent == 100
    assert job.error_message is None


def test_final_internal_json_is_saved_and_valid(session, tmp_path):
    course = _make_course(session)

    job = run_generation(session, course.id)

    save_log = next(entry for entry in job.log_json if entry["step"] == "save_internal_json")
    saved_path = save_log["path"]
    data = json.loads(open(saved_path, encoding="utf-8").read())

    assert data["title"] == course.title
    assert "full_text" in data and data["full_text"]
    assert len(data["modules"]) == FakeProvider.DEFAULT_MODULE_COUNT


def test_run_generation_creates_course_version_and_docx(session):
    course = _make_course(session)

    job = run_generation(session, course.id)

    versions = course_versions.list(session, course_id=course.id)
    assert len(versions) == 1
    assert versions[0].version_number == 1
    assert job.output_docx_path == versions[0].output_docx_path

    docx_path = versions[0].output_docx_path
    assert docx_path.endswith("course_v1.docx")

    document = Document(docx_path)
    texts = [p.text for p in document.paragraphs]
    assert course.title in texts
    # Every module title from the fake map should appear as a heading.
    assert any("Module 1" in t for t in texts)
    assert any("Module 2" in t for t in texts)


def test_regenerating_a_course_creates_a_new_version(session):
    course = _make_course(session)

    run_generation(session, course.id)
    run_generation(session, course.id)

    versions = course_versions.list(session, course_id=course.id)
    assert sorted(v.version_number for v in versions) == [1, 2]


def test_logs_are_present_and_short(session):
    course = _make_course(session)

    job = run_generation(session, course.id)

    assert len(job.log_json) > 0
    steps = [entry["step"] for entry in job.log_json]
    assert "load_context" in steps
    assert "build_map" in steps
    assert "final_review" in steps
    assert "complete" in steps
    # "Keep logs short": no single entry should balloon into a huge blob
    # (e.g. never a raw script_text sneaking into a log entry).
    for entry in job.log_json:
        assert len(json.dumps(entry)) < 300


def test_final_review_pass_skips_ai_rebuild(session):
    course = _make_course(session)

    job = run_generation(session, course.id)

    rebuild_log = next(e for e in job.log_json if e["step"] == "rebuild_final_course")
    assert rebuild_log["triggered"] is False


def test_active_admin_knowledge_is_loaded_into_context(session):
    course = _make_course(session)
    admin_knowledge_items.create(
        session,
        key="rukn-core",
        title="Core rules",
        item_type=ItemType.MARKDOWN,
        content_text="Some rule text",
        is_active=True,
    )
    admin_knowledge_items.create(
        session,
        key="inactive-rule",
        title="Old rules",
        item_type=ItemType.MARKDOWN,
        content_text="Should not be loaded",
        is_active=False,
    )

    job = run_generation(session, course.id)

    load_log = next(e for e in job.log_json if e["step"] == "load_context")
    assert load_log["rules"] == 1


def test_only_usable_sources_are_counted(session):
    course = _make_course(session)
    course_sources.create(
        session,
        course_id=course.id,
        source_category=SourceCategory.SCIENTIFIC_REFERENCE,
        priority=Priority.HIGH,
        status="ready",
        extracted_text="Usable extracted text.",
    )
    course_sources.create(
        session,
        course_id=course.id,
        source_category=SourceCategory.SCIENTIFIC_REFERENCE,
        priority=Priority.MEDIUM,
        status="extraction_blocked",
        extracted_text=None,
    )

    job = run_generation(session, course.id)

    load_log = next(e for e in job.log_json if e["step"] == "load_context")
    assert load_log["sources"] == 1


def test_map_gets_full_text_for_short_source(session):
    course = _make_course(session)
    short_text = "Short note about Excel VLOOKUP."
    _make_source_with_analysis(session, course.id, short_text)
    provider = RecordingProvider()

    run_generation(session, course.id, provider=provider)

    assert provider.map_calls[0].sources[0].text == short_text


def test_map_gets_summary_not_full_text_for_long_source(session):
    course = _make_course(session)
    long_text = (
        "# Excel Basics\n"
        + ("Opening Excel and understanding the ribbon interface. " * 40)
        + "\n\n# Advanced Formulas\n"
        + ("VLOOKUP and SUMIF are powerful budgeting formulas. " * 40)
    )
    assert len(long_text) > SHORT_SOURCE_MAX_CHARS
    _make_source_with_analysis(session, course.id, long_text)
    provider = RecordingProvider()

    run_generation(session, course.id, provider=provider)

    excerpt_text = provider.map_calls[0].sources[0].text
    assert excerpt_text
    assert len(excerpt_text) < len(long_text)


def test_reel_gets_full_text_for_short_source(session):
    course = _make_course(session)
    short_text = "Short note about Excel VLOOKUP."
    _make_source_with_analysis(session, course.id, short_text)
    provider = RecordingProvider()

    run_generation(session, course.id, provider=provider)

    # Not necessarily exactly one call per reel: FakeProvider's templated
    # openings can legitimately trigger the local opening_checker's
    # bounded retries (see app/validators/opening_checker.py) - that's
    # correct behavior, not a bug. What matters here is every attempt got
    # the full short source text.
    assert len(provider.reel_calls) >= FakeProvider.DEFAULT_MODULE_COUNT * FakeProvider.DEFAULT_REELS_PER_MODULE
    for reel_input in provider.reel_calls:
        assert reel_input.sources[0].text == short_text


def test_reel_never_gets_full_text_of_a_long_source(session):
    course = _make_course(session)
    long_text = (
        "# Excel Basics\n"
        + ("Opening Excel and understanding the ribbon interface. " * 40)
        + "\n\n# Advanced Formulas\n"
        + ("VLOOKUP and SUMIF are powerful budgeting formulas. " * 40)
    )
    _make_source_with_analysis(session, course.id, long_text)
    provider = RecordingProvider()

    run_generation(session, course.id, provider=provider)

    assert len(provider.reel_calls) > 0
    for reel_input in provider.reel_calls:
        excerpt_text = reel_input.sources[0].text
        assert excerpt_text != long_text
        assert len(excerpt_text) < len(long_text)


def test_local_validators_run_before_ai_review_and_can_skip_it(session):
    """FakeProvider's templated openings legitimately trip the local
    opening_checker (see the comment on test_reel_gets_full_text_for_short_source
    above) - so across a full run, at least one reel should be caught
    locally, meaning the AI review call was skipped for that attempt."""
    course = _make_course(session)
    provider = RecordingProvider()

    job = run_generation(session, course.id, provider=provider)

    reel_logs = [e for e in job.log_json if e["step"] == "reel"]
    assert any(e["caught_locally"] for e in reel_logs)
    # Every locally-caught attempt means one fewer AI review call than
    # total write attempts.
    assert provider.review_single_reel_calls < len(provider.reel_calls)


def test_forbidden_phrase_short_circuits_ai_review_entirely(session):
    """Direct, precise test of the short-circuit: a provider whose reel
    always contains a forbidden phrase should never reach the AI reviewer."""
    from app.generation.orchestrator import _write_and_review_reel
    from app.schemas.generation import CourseMap, GeneratedReel, ModulePlan, ReelPlan, ReviewStatus

    class AlwaysForbiddenProvider(FakeProvider):
        def __init__(self):
            super().__init__()
            self.review_calls = 0

        def write_single_reel(self, input):  # noqa: A002
            return GeneratedReel(
                reel_id=input.reel.reel_id,
                module_id=input.module.module_id,
                title=input.reel.title,
                script_text="في الريل ده هنشرح حاجة مهمة جدا اليوم عن الموضوع ده وهنكمل كلامنا",
                used_ideas=["idea"],
                used_examples=["example"],
                self_check_status=ReviewStatus.PASS,
            )

        def review_single_reel(self, input):  # noqa: A002
            self.review_calls += 1
            return super().review_single_reel(input)

    module = ModulePlan(module_id="m1", title="Module", purpose="p", reels=[])
    reel_plan = ReelPlan(reel_id="m1-r1", title="Reel", purpose="p", estimated_length="30s")
    course_map = CourseMap(course_title="Course", main_thread="thread", modules=[module])
    provider = AlwaysForbiddenProvider()
    rules_context = {
        "rukn_forbidden_phrases": json.dumps(
            {"phrases": [{"phrase": "في الريل ده", "severity": "high"}]}
        )
    }

    generated, attempts, caught_locally = _write_and_review_reel(
        provider=provider,
        course_map=course_map,
        module=module,
        reel_plan=reel_plan,
        prior_reels=[],
        all_reels_so_far=[],
        sources=[],
        rules_context=rules_context,
    )

    assert caught_locally is True
    assert provider.review_calls == 0
    assert attempts == 3  # 1 initial + 2 retries (MAX_REEL_REWRITE_ATTEMPTS)
    assert generated.script_text  # still returns the (flagged) reel, not None


def test_stage_history_includes_reading_sources_and_reviewing_repetition(session, monkeypatch):
    """Verifies the transient stage values actually get set at some point
    during the run (they're overwritten by later flushes, so we have to
    watch every generation_jobs.update call, not just the final job)."""
    import app.generation.orchestrator as orchestrator_module

    stage_history: list[str] = []
    original_update = orchestrator_module.generation_jobs.update

    def spy_update(session_, id_, **fields):
        if "current_stage" in fields:
            stage_history.append(fields["current_stage"])
        return original_update(session_, id_, **fields)

    monkeypatch.setattr(orchestrator_module.generation_jobs, "update", spy_update)

    course = _make_course(session)
    run_generation(session, course.id)

    assert "reading_sources" in stage_history
    assert "building_map" in stage_history
    assert "generating" in stage_history
    assert "reviewing_repetition" in stage_history
    assert "reviewing" in stage_history
    assert "exporting" in stage_history
    assert "done" in stage_history


def test_summary_text_is_meaningful(session):
    course = _make_course(session)

    run_generation(session, course.id)

    version = course_versions.list(session, course_id=course.id)[0]
    assert course.title in version.summary_text
    assert "module" in version.summary_text.lower()
    assert "reel" in version.summary_text.lower()


def test_report_text_only_populated_for_full_report_level(session):
    plain_course = _make_course(session, explanation_level=ExplanationLevel.FINAL_ONLY)
    run_generation(session, plain_course.id)
    plain_version = course_versions.list(session, course_id=plain_course.id)[0]
    assert plain_version.report_text is None

    full_report_course = _make_course(
        session, title="Full Report Course", explanation_level=ExplanationLevel.FULL_REPORT
    )
    run_generation(session, full_report_course.id)
    full_version = course_versions.list(session, course_id=full_report_course.id)[0]
    assert full_version.report_text is not None
    assert "Course:" in full_version.report_text
    assert "Review checkpoints run" in full_version.report_text


def test_two_modules_review_runs_for_even_module_count(session):
    course = _make_course(session)

    job = run_generation(session, course.id)

    two_module_logs = [e for e in job.log_json if e["step"] == "review_2modules"]
    assert len(two_module_logs) == 1
    assert "skipped" not in two_module_logs[0]
    assert two_module_logs[0]["ids"] == ["m1", "m2"]


def test_two_modules_review_skips_unpaired_trailing_module(session):
    class ThreeModuleProvider(FakeProvider):
        DEFAULT_MODULE_COUNT = 3

    course = _make_course(session)

    job = run_generation(session, course.id, provider=ThreeModuleProvider())

    two_module_logs = [e for e in job.log_json if e["step"] == "review_2modules"]
    assert len(two_module_logs) == 2
    assert two_module_logs[0]["ids"] == ["m1", "m2"]
    assert two_module_logs[1].get("skipped") == "unpaired trailing module"


def test_five_reel_review_runs_once_for_six_reel_course(session):
    course = _make_course(session)

    job = run_generation(session, course.id)

    five_reel_logs = [e for e in job.log_json if e["step"] == "review_5reels"]
    # 2 modules x 3 reels = 6 reels total -> exactly one trigger at reel 5.
    assert len(five_reel_logs) == 1


def test_course_not_found_raises(session):
    with pytest.raises(ValueError):
        run_generation(session, course_id=999999)


def test_source_style_and_jargon_never_leak_into_exported_docx(session):
    """Source Authority Firewall, end to end: a scientific_reference source
    with distinctive academic jargon and a flow_reference source with a
    distinctive catchphrase must never appear verbatim in the exported
    DOCX, and the DOCX must still follow the standard teleprompter
    structure regardless of the sources' own formatting/tone. This
    currently passes largely because FakeProvider never consumes source
    text into its placeholder script - that's fine, it's a forward-looking
    regression guard for when a real provider is wired in."""
    course = _make_course(session)

    jargon_phrase = "epistemological paradigmatic heuristic substantiation"
    academic_text = (
        f"In accordance with the {jargon_phrase} framework established by prior "
        "scholarship, the present treatise shall proceed to enumerate the "
        "constituent formulae. " * 10
    )
    _make_source_with_analysis(
        session, course.id, academic_text, category=SourceCategory.SCIENTIFIC_REFERENCE
    )

    catchphrase = "يلا يا نجم يلا"
    flow_text = (
        f"{catchphrase} خد بالك من الخطوة الأولى كويس وابدأ فعليا من غير "
        f"تفكير كتير. لو الألوان مش واضحة هتضيع وقتك. {catchphrase} جرب دلوقتي."
    )
    _make_source_with_analysis(
        session, course.id, flow_text, category=SourceCategory.FLOW_REFERENCE
    )

    job = run_generation(session, course.id)

    assert job.status == JobStatus.COMPLETED
    document = Document(job.output_docx_path)
    paragraph_texts = [p.text for p in document.paragraphs]
    full_text = "\n".join(paragraph_texts)

    assert jargon_phrase not in full_text
    assert catchphrase not in full_text

    # Standard teleprompter structure still holds regardless of the
    # sources' own formatting/tone.
    assert any(t.startswith("Module 1 —") for t in paragraph_texts)
    assert any(t.startswith("Lesson 1 —") for t in paragraph_texts)


def test_provider_error_with_no_saved_work_marks_job_failed_with_clean_message(session):
    """Nothing usable was saved yet (fails before the course map even
    finishes) - the job must end FAILED, with a short, clean error_message
    that never contains the raw exception text (that stays in the
    internal-only log_json entry, never in what the API returns)."""

    class BrokenProvider(FakeProvider):
        def build_course_map(self, input):  # noqa: A002 - matches AIProvider's signature
            raise RuntimeError("simulated rate limit: HTTP 429 too many requests")

    course = _make_course(session)

    job = run_generation(session, course.id, provider=BrokenProvider())

    assert job.status == JobStatus.FAILED
    assert job.current_stage == "failed"
    assert job.course_map_json is None
    assert job.completed_reels_json == []
    assert job.partial_docx_path is None
    assert job.error_category == "rate_limit"
    assert job.error_message == ERROR_CATEGORY_MESSAGES_NOTHING_SAVED["rate_limit"]
    assert "simulated rate limit" not in job.error_message
    assert "429" not in job.error_message

    error_logs = [e for e in job.log_json if e["step"] == "error"]
    assert len(error_logs) == 1
    # The raw text is allowed in the internal-only log, just never in the
    # user-facing error_message.
    assert "simulated rate limit" in error_logs[0]["message"]


def test_provider_error_with_saved_work_marks_job_partial_and_saves_partial_docx(session):
    """A course map and some completed reels already exist when the
    failure happens - the job must end PARTIAL (not FAILED), with a clean
    category-specific error_message and a real partial DOCX on disk."""
    course = _make_course(session)
    provider = FailAfterNReelsProvider(
        fail_after=2, message="Service unavailable: 503 - please retry later"
    )

    job = run_generation(session, course.id, provider=provider)

    assert job.status == JobStatus.PARTIAL
    assert job.current_stage == "partial"
    assert job.course_map_json is not None
    assert len(job.completed_reels_json) == 2
    assert job.completed_reels_count == 2
    assert job.completed_modules_count == 0  # module 1's review never ran
    assert job.last_completed_step == "reel:m1-r2"
    assert job.error_category == "provider_unavailable"
    assert job.error_message == ERROR_CATEGORY_MESSAGES["provider_unavailable"]
    assert "503" not in job.error_message

    assert job.partial_docx_path is not None
    assert Path(job.partial_docx_path).exists()


def test_realistic_anthropic_shaped_rate_limit_error_ends_job_partial(session):
    """Closes the gap between the generic reliability-pass tests above
    (plain RuntimeError) and what a real Anthropic SDK failure actually
    looks like: a distinctly-named exception class with 429/rate-limit
    wording. The job must still end PARTIAL with the already-completed
    reels intact."""
    course = _make_course(session)
    provider = RealisticRateLimitProvider(fail_after=2)

    job = run_generation(session, course.id, provider=provider)

    assert job.status == JobStatus.PARTIAL
    assert job.current_stage == "partial"
    assert job.error_category == "rate_limit"
    assert job.course_map_json is not None
    assert len(job.completed_reels_json) == 2
    assert job.completed_reels_count == 2

    reloaded = generation_jobs.get(session, job.id)
    assert reloaded.status == JobStatus.PARTIAL
    assert len(reloaded.completed_reels_json) == 2


def test_partial_run_does_not_lose_or_overwrite_already_completed_work(session):
    """Re-fetch the job row from the DB after the failure (not just the
    in-memory return value) to confirm nothing got wiped."""
    course = _make_course(session)
    provider = FailAfterNReelsProvider(fail_after=2)

    job = run_generation(session, course.id, provider=provider)
    assert job.status == JobStatus.PARTIAL

    reloaded = generation_jobs.get(session, job.id)
    assert reloaded is not None
    assert reloaded.status == JobStatus.PARTIAL
    assert reloaded.course_map_json is not None
    assert len(reloaded.completed_reels_json) == 2
    assert reloaded.completed_reels_json == job.completed_reels_json


def test_regenerate_from_scratch_is_independent_of_an_existing_partial_job(session):
    """`run_generation` always creates a brand-new `GenerationJob` row -
    starting fresh must not be blocked by, or corrupt, an existing
    partial job's saved state."""
    course = _make_course(session)
    partial_job = run_generation(session, course.id, provider=FailAfterNReelsProvider(fail_after=2))
    assert partial_job.status == JobStatus.PARTIAL

    fresh_job = run_generation(session, course.id)

    assert fresh_job.id != partial_job.id
    assert fresh_job.status == JobStatus.COMPLETED

    reloaded_partial = generation_jobs.get(session, partial_job.id)
    assert reloaded_partial.status == JobStatus.PARTIAL
    assert len(reloaded_partial.completed_reels_json) == 2
    assert reloaded_partial.course_map_json is not None

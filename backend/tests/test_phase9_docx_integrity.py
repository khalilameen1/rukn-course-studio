"""Phase 9 Teleprompter DOCX package, direction, order, and render gates."""

import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.generation.quality.context_snapshot import fingerprint_value
from app.generation.teleprompter_checks import module_lesson_structure_present
from app.models.enums import AddressForm
from app.schemas.generation import (
    CourseThesis,
    FinalCourse,
    FinalModule,
    FinalReel,
    ModuleProject,
)
from app.services.docx_export import render_final_course_docx
from app.services.docx_verification import (
    DocxVerificationError,
    assert_final_course_ready_for_docx,
    verify_docx_archive,
)


def _course() -> FinalCourse:
    thesis = CourseThesis(
        final_student_outcome="تنفيذ قرار بصري واضح",
        audience_and_starting_level="مصمم مبتدئ",
        practical_deliverable="تصميم واحد معدل",
        student_language="ar",
        spoken_variety="egyptian_colloquial",
        address_form=AddressForm.MASCULINE,
    )
    return FinalCourse(
        title="استوديو القرار البصري",
        full_text="نص داخلي لا يدخل الملف",
        thesis=thesis,
        modules=[
            FinalModule(
                module_id="m1",
                title="الهرمية",
                reels=[
                    FinalReel(
                        reel_id="m1-r1",
                        title="وزن العنوان في Figma",
                        script_text=(
                            "العنوان الأهم ياخد وزن أوضح من باقي العناصر\n"
                            "جرّب النسخة على Figma وراقب أول نقطة تمسكها عينك"
                        ),
                        spoken_beats=[
                            "العنوان الأهم ياخد وزن أوضح من باقي العناصر",
                            "جرّب النسخة على Figma وراقب أول نقطة تمسكها عينك",
                        ],
                    )
                ],
                module_project=ModuleProject(
                    name="مشروع الموديول الأول",
                    brief="عدّل تصميم واحد وخلي ترتيب القراءة واضح من أول نظرة",
                    inputs_or_files=["INTERNAL asset list"],
                    deliverable_shape="INTERNAL deliverable metadata",
                    pass_criteria=["INTERNAL rubric"],
                    skills_tested=["INTERNAL skill ledger"],
                ),
            ),
            FinalModule(
                module_id="m2",
                title="التباين",
                reels=[
                    FinalReel(
                        reel_id="m2-r1",
                        title="فرق الإضاءة",
                        script_text=(
                            "قارن الزرار بالخلفية قبل ما تزود أي لون\n"
                            "لو العين مسكت الزرار بسرعة يبقى القرار اشتغل"
                        ),
                        spoken_beats=[
                            "قارن الزرار بالخلفية قبل ما تزود أي لون",
                            "لو العين مسكت الزرار بسرعة يبقى القرار اشتغل",
                        ],
                    )
                ],
                module_project=ModuleProject(
                    name="مشروع الموديول التاني",
                    brief="اختبر نسختين وسجل أي نسخة اتقرت أسرع",
                ),
            ),
        ],
        graduation_project=ModuleProject(
            name="مشروع التخرج",
            brief="سلّم تصميم نهائي واشرح سبب كل قرار بصري فيه",
            pass_criteria=["INTERNAL capstone rubric"],
        ),
    )


def test_preflight_blocks_mutation_metadata_stale_beats_and_dense_layout():
    course = _course()
    dirty = course.model_copy(deep=True)
    dirty.modules[0].reels[0] = dirty.modules[0].reels[0].model_copy(
        update={
            "script_text": (
                "Hook: ملاحظة داخلية فيها علامات ترقيم ثم "
                + " ".join(f"كلمة{i}" for i in range(30))
            ),
            "quality_status": "needs_review",
        }
    )
    with pytest.raises(DocxVerificationError) as caught:
        assert_final_course_ready_for_docx(dirty)
    message = str(caught.value)
    assert "quality_status=needs_review" in message
    assert "spoken_beats differ" in message
    assert "punctuation mutation" in message
    assert "metadata" in message
    assert "dense spoken line" in message


def test_docx_zip_xml_headings_rtl_ltr_projects_and_exact_text(tmp_path):
    course = _course()
    assert_final_course_ready_for_docx(course)
    path = tmp_path / "phase9.docx"
    render_final_course_docx(course).save(path)

    report = verify_docx_archive(path, course)
    assert report.ok, report.errors
    assert report.accepted_text_sha256 == report.docx_text_sha256
    assert len(report.accepted_text_sha256) == 64

    with zipfile.ZipFile(path) as archive:
        assert archive.testzip() is None
        assert "word/document.xml" in archive.namelist()
        assert "word/styles.xml" in archive.namelist()
        xml = archive.read("word/document.xml").decode("utf-8")
        assert "INTERNAL asset list" not in xml
        assert "INTERNAL rubric" not in xml
        assert "INTERNAL capstone rubric" not in xml
        assert "w:bidi" in xml
        assert 'w:rtl w:val="0"' in xml
        assert "w:keepNext" in xml

    reopened = Document(path)
    visible = [paragraph.text for paragraph in reopened.paragraphs if paragraph.text]
    assert module_lesson_structure_present("\n".join(visible))
    assert visible[0] == "استوديو القرار البصري"
    assert "الموديول 1 الهرمية" in visible
    assert "الريل 1 وزن العنوان في Figma" in visible
    assert visible.count("مشروع التخرج") == 1
    assert visible.index("مشروع الموديول الأول") < visible.index("الموديول 2 التباين")
    assert visible[-2:] == [
        "مشروع التخرج",
        "سلّم تصميم نهائي واشرح سبب كل قرار بصري فيه",
    ]

    for paragraph in reopened.paragraphs:
        if any("\u0600" <= char <= "\u06ff" for char in paragraph.text):
            assert paragraph._p.pPr is not None  # noqa: SLF001
            assert paragraph._p.pPr.find(qn("w:bidi")) is not None  # noqa: SLF001

    mixed = next(paragraph for paragraph in reopened.paragraphs if "Figma" in paragraph.text)
    latin_runs = [run for run in mixed.runs if "Figma" in run.text]
    assert latin_runs
    assert all(
        run._r.rPr.find(qn("w:rtl")).get(qn("w:val")) == "0"  # noqa: SLF001
        for run in latin_runs
    )


def test_docx_geometry_matches_teleprompter_preset_override():
    document = render_final_course_docx(_course())
    section = document.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert round(section.left_margin.inches, 2) == 1.0
    assert round(section.right_margin.inches, 2) == 1.0
    assert document.styles["Normal"].font.name == "Arial"
    assert document.styles["Normal"].font.size.pt == 14
    assert document.styles["Normal"].paragraph_format.line_spacing == 1.5
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style_p_pr = document.styles[style_name].element.get_or_add_pPr()
        assert style_p_pr.find(qn("w:pBdr")) is None
    assert not document.tables


def test_visible_fingerprint_changes_when_accepted_spoken_text_changes(tmp_path):
    first = _course()
    first_path = tmp_path / "first.docx"
    render_final_course_docx(first).save(first_path)
    first_report = verify_docx_archive(first_path, first)
    assert first_report.ok

    second = _course().model_copy(deep=True)
    changed = second.modules[0].reels[0].model_copy(
        update={
            "script_text": (
                "العنوان الأهم ياخد وزن أقوى من باقي العناصر\n"
                "جرّب النسخة على Figma وراقب أول نقطة تمسكها عينك"
            ),
            "spoken_beats": [
                "العنوان الأهم ياخد وزن أقوى من باقي العناصر",
                "جرّب النسخة على Figma وراقب أول نقطة تمسكها عينك",
            ],
        }
    )
    second.modules[0].reels[0] = changed
    second_path = tmp_path / "second.docx"
    render_final_course_docx(second).save(second_path)
    second_report = verify_docx_archive(second_path, second)
    assert second_report.ok
    assert first_report.accepted_text_sha256 != second_report.accepted_text_sha256
    assert fingerprint_value(first_report.accepted_text_sha256) != fingerprint_value(
        second_report.accepted_text_sha256
    )

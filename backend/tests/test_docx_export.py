"""Tests for app/services/docx_export.py - the teleprompter DOCX contract."""

from docx import Document

from app.generation.teleprompter_checks import TELEPROMPTER_FORBIDDEN_SUBSTRINGS
from app.schemas.generation import (
    CourseMap,
    FinalCourse,
    FinalModule,
    FinalReel,
    GeneratedReel,
    ModulePlan,
    ModuleProject,
    ReelPlan,
    ReviewStatus,
)
from app.services.docx_export import (
    PARTIAL_DRAFT_NOTE,
    build_partial_course_from_job,
    export_final_course_to_docx,
    export_partial_course_to_docx,
    next_version_number,
    render_final_course_docx,
    render_partial_course_docx,
)

# Substrings that must never appear anywhere in the rendered DOCX text, per
# the rukn_teleprompter_docx_contract admin knowledge item - the DOCX must
# hide every internal-pipeline artifact (review notes, validation notes,
# quality checks, etc.), never show credit/methodology text, and never
# address the lecturer with meta-instructions instead of actual lines to
# say. Sourced from app/generation/teleprompter_checks.py (shared with
# app/generation/output_scoring.py and backend/tests/golden/) rather than
# duplicated here, so this list can't silently drift between the two.
FORBIDDEN_SUBSTRINGS = list(TELEPROMPTER_FORBIDDEN_SUBSTRINGS)


def _sample_final_course() -> FinalCourse:
    return FinalCourse(
        title="Intro to Excel Formulas",
        modules=[
            FinalModule(
                module_id="m1",
                title="Getting Started",
                bridge_project="INTERNAL BRIDGE — never in DOCX",
                module_project=ModuleProject(
                    name="مشروع موديول 1",
                    brief="Build a starter budget sheet",
                    deliverable_shape="ملف ميزانية",
                    pass_criteria=["يبني الجدول"],
                    skills_tested=["excel-basics"],
                ),
                reels=[
                    FinalReel(reel_id="m1-r1", title="Opening Excel", script_text="Line one.\nLine two."),
                    FinalReel(reel_id="m1-r2", title="Basic Formulas", script_text="Type an equals sign."),
                ],
            ),
            FinalModule(
                module_id="m2",
                title="Budgets",
                bridge_project=None,
                module_project=ModuleProject(
                    name="مشروع موديول 2",
                    brief="Apply SUM on a real budget file",
                    deliverable_shape="ملف نهائي",
                    pass_criteria=["يستخدم SUM"],
                    skills_tested=["sum"],
                ),
                reels=[
                    FinalReel(reel_id="m2-r1", title="Totals", script_text="Use SUM to add a column."),
                ],
            ),
        ],
        full_text="irrelevant - the exporter renders from `modules`, not this",
        graduation_project=ModuleProject(
            name="مشروع التخرج",
            brief="Complete workbook",
            deliverable_shape="ملف كامل",
            pass_criteria=["يغطي المهارات"],
            skills_tested=["capstone"],
        ),
    )


def _headings(document) -> list[tuple[str, str]]:
    """[(style_name, text), ...] for every heading paragraph in the doc."""
    return [
        (p.style.name, p.text)
        for p in document.paragraphs
        if p.style.name.startswith("Heading") or p.style.name == "Title"
    ]


def _all_text(document) -> str:
    return "\n".join(p.text for p in document.paragraphs)


def _sample_partial_course_map() -> CourseMap:
    """Module 2 (and module 1's second reel) intentionally have no
    completed reel behind them - simulates a run that stopped partway
    through module 1."""
    return CourseMap(
        course_title="Intro to Excel Formulas",
        main_thread="thread",
        modules=[
            ModulePlan(
                module_id="m1",
                title="Getting Started",
                purpose="p",
                bridge_project=None,
                module_project=ModuleProject(
                    name="مشروع موديول 1",
                    brief="Build a starter budget sheet",
                ),
                reels=[
                    ReelPlan(reel_id="m1-r1", title="Opening Excel", purpose="p", estimated_length="30s"),
                    ReelPlan(reel_id="m1-r2", title="Basic Formulas", purpose="p", estimated_length="30s"),
                ],
            ),
            ModulePlan(
                module_id="m2",
                title="Budgets",
                purpose="p",
                reels=[ReelPlan(reel_id="m2-r1", title="Totals", purpose="p", estimated_length="30s")],
            ),
        ],
    )


def _sample_completed_reels() -> list[GeneratedReel]:
    """Only module 1's first reel actually completed before the run
    stopped - `m1-r2` and all of module 2 never got written."""
    return [
        GeneratedReel(
            reel_id="m1-r1",
            module_id="m1",
            title="Opening Excel",
            script_text="Line one.\nLine two.",
            self_check_status=ReviewStatus.PASS,
        )
    ]


def test_render_includes_course_title_as_top_level_heading():
    document = render_final_course_docx(_sample_final_course())

    headings = _headings(document)
    assert headings[0] == ("Title", "Intro to Excel Formulas")


def test_render_numbers_modules_and_lessons_in_order():
    document = render_final_course_docx(_sample_final_course())

    headings = [text for style, text in _headings(document) if style != "Title"]
    # Module projects render as H2 between modules (not Lesson N).
    assert headings[0] == "Module 1 — Getting Started"
    assert "Lesson 1 — Opening Excel" in headings
    assert "Lesson 2 — Basic Formulas" in headings
    assert "Module 2 — Budgets" in headings
    assert "Lesson 1 — Totals" in headings
    assert any("مشروع" in h or "Build a starter" in h for h in headings)


def test_render_includes_script_text_under_each_lesson():
    document = render_final_course_docx(_sample_final_course())

    body_texts = [p.text for p in document.paragraphs if not p.style.name.startswith("Heading")]
    # Punctuation stripped from spoken body.
    assert "Line one" in body_texts
    assert "Line two" in body_texts
    assert "Type an equals sign" in body_texts
    assert "Use SUM to add a column" in body_texts


def test_render_includes_module_projects_not_as_lessons():
    """Module projects appear between modules; never as Lesson N."""
    document = render_final_course_docx(_sample_final_course())

    headings = [text for style, text in _headings(document) if style != "Title"]
    lesson_headings = [h for h in headings if h.startswith("Lesson ")]
    assert lesson_headings == [
        "Lesson 1 — Opening Excel",
        "Lesson 2 — Basic Formulas",
        "Lesson 1 — Totals",
    ]
    body = _all_text(document)
    assert "Build a starter budget sheet" in body


def test_render_excludes_internal_pipeline_language():
    document = render_final_course_docx(_sample_final_course())

    text = _all_text(document).lower()
    for forbidden in FORBIDDEN_SUBSTRINGS:
        assert forbidden not in text


def test_next_version_number():
    assert next_version_number([]) == 1
    assert next_version_number([1]) == 2
    assert next_version_number([1, 2, 3]) == 4
    assert next_version_number([5, 1, 3]) == 6


def test_export_final_course_to_docx_writes_real_openable_file(tmp_path, monkeypatch):
    import app.services.docx_export as docx_export_module

    monkeypatch.setattr(docx_export_module.settings, "storage_outputs_dir", tmp_path)

    path = export_final_course_to_docx(_sample_final_course(), course_id=42, version_number=3)

    assert path == tmp_path / "42" / "course_v3.docx"
    assert path.exists()

    reopened = Document(str(path))
    texts = [p.text for p in reopened.paragraphs]
    assert "Intro to Excel Formulas" in texts
    assert "Module 1 — Getting Started" in texts


def test_final_docx_never_contains_partial_draft_note():
    """Regression guard: a normal successful run's DOCX must never contain
    the partial-draft marker - that paragraph only ever comes from
    `render_partial_course_docx`, never `render_final_course_docx`."""
    document = render_final_course_docx(_sample_final_course())

    texts = [p.text for p in document.paragraphs]
    assert PARTIAL_DRAFT_NOTE not in texts
    assert "partial draft" not in _all_text(document).lower()


def test_build_partial_course_from_job_skips_modules_with_no_completed_reels():
    course_map = _sample_partial_course_map()
    reels = _sample_completed_reels()

    partial_course = build_partial_course_from_job(
        course_map.model_dump(mode="json"), [r.model_dump(mode="json") for r in reels]
    )

    assert len(partial_course.modules) == 1
    assert partial_course.modules[0].module_id == "m1"
    assert [r.reel_id for r in partial_course.modules[0].reels] == ["m1-r1"]


def test_render_partial_course_docx_includes_note_and_only_completed_content():
    course_map = _sample_partial_course_map()
    reels = _sample_completed_reels()
    partial_course = build_partial_course_from_job(
        course_map.model_dump(mode="json"), [r.model_dump(mode="json") for r in reels]
    )

    document = render_partial_course_docx(partial_course)
    texts = [p.text for p in document.paragraphs]

    assert texts[0] == PARTIAL_DRAFT_NOTE
    assert "Intro to Excel Formulas" in texts
    assert any("Getting Started" in t for t in texts)
    # Not-yet-completed module/reel must not appear.
    assert not any("Budgets" in t for t in texts)
    assert not any("Basic Formulas" in t for t in texts)


def test_render_partial_course_docx_excludes_internal_pipeline_language():
    course_map = _sample_partial_course_map()
    reels = _sample_completed_reels()
    partial_course = build_partial_course_from_job(
        course_map.model_dump(mode="json"), [r.model_dump(mode="json") for r in reels]
    )

    document = render_partial_course_docx(partial_course)
    text = _all_text(document).lower()
    for forbidden in FORBIDDEN_SUBSTRINGS:
        assert forbidden not in text


def test_export_partial_course_to_docx_uses_a_distinct_naming_pattern(tmp_path, monkeypatch):
    import app.services.docx_export as docx_export_module

    monkeypatch.setattr(docx_export_module.settings, "storage_outputs_dir", tmp_path)

    course_map = _sample_partial_course_map()
    reels = _sample_completed_reels()
    partial_course = build_partial_course_from_job(
        course_map.model_dump(mode="json"), [r.model_dump(mode="json") for r in reels]
    )

    path = export_partial_course_to_docx(partial_course, course_id=42, job_id=7)

    # Distinct from a real "course_v{n}.docx" version - never confusable
    # with a completed version, never picked up by next_version_number or
    # version listings (both only ever look at CourseVersion DB rows).
    assert path == tmp_path / "42" / "partial_job_7.docx"
    assert path.exists()

    reopened = Document(str(path))
    assert reopened.paragraphs[0].text == PARTIAL_DRAFT_NOTE

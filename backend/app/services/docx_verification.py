"""Fail-closed DOCX package, text, direction, and rendered-page verification."""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
import zipfile
import hashlib
from dataclasses import dataclass, field
from pathlib import Path

import fitz
from docx import Document
from docx.oxml.ns import qn

from app.generation.contracts.spoken_final_master import (
    beats_to_plain_script,
    strip_punctuation_from_spoken_body,
    validate_spoken_export_text,
)
from app.generation.teleprompter_checks import find_forbidden_substrings
from app.schemas.generation import FinalCourse

_REQUIRED_DOCX_PARTS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "word/document.xml",
    "word/styles.xml",
}
_ARABIC = re.compile(r"[\u0600-\u06FF]")
_ONE_WORD_EXCEPTIONS = {
    "أيوه",
    "لأ",
    "لا",
    "صح",
    "غلط",
    "تمام",
    "ليه",
    "yes",
    "no",
    "why",
}


class DocxVerificationError(RuntimeError):
    """The export is structurally, textually, directionally, or visually unsafe."""


@dataclass
class DocxVerificationReport:
    errors: list[str] = field(default_factory=list)
    page_count: int = 0
    rendered_page_paths: list[Path] = field(default_factory=list)
    accepted_text_sha256: str = ""
    docx_text_sha256: str = ""

    @property
    def ok(self) -> bool:
        return not self.errors

    def raise_if_invalid(self) -> None:
        if self.errors:
            raise DocxVerificationError("; ".join(self.errors))


def _normalized(text: str) -> str:
    return "\n".join(
        line.strip()
        for line in (text or "").replace("\r", "").splitlines()
        if line.strip()
    )


def _sha256(text: str) -> str:
    return hashlib.sha256((text or "").encode("utf-8")).hexdigest()


def _course_is_arabic(course: FinalCourse) -> bool:
    if course.thesis is not None:
        return str(course.thesis.student_language or "").lower().startswith("ar")
    return bool(_ARABIC.search(course.title or ""))


def module_heading(course: FinalCourse, index: int, title: str) -> str:
    return f"الموديول {index} {title}" if _course_is_arabic(course) else f"Module {index} — {title}"


def lesson_heading(course: FinalCourse, index: int, title: str) -> str:
    return f"الريل {index} {title}" if _course_is_arabic(course) else f"Lesson {index} — {title}"


def assert_final_course_ready_for_docx(course: FinalCourse) -> None:
    """Reject any text mutation, metadata, or visually unsafe line before save."""
    errors: list[str] = []
    if not (course.title or "").strip():
        errors.append("course title is empty")
    if not course.modules:
        errors.append("course has no modules")

    for module in course.modules:
        if not module.reels:
            errors.append(f"module {module.module_id} has no lessons")
        for reel in module.reels:
            text = reel.script_text or ""
            if reel.quality_status != "pass":
                errors.append(f"{reel.reel_id}: quality_status={reel.quality_status}")
            if not text.strip():
                errors.append(f"{reel.reel_id}: empty spoken text")
                continue
            if reel.spoken_beats:
                from_beats = beats_to_plain_script(reel.spoken_beats)
                if _normalized(from_beats) != _normalized(text):
                    errors.append(f"{reel.reel_id}: spoken_beats differ from accepted script")
            if strip_punctuation_from_spoken_body(text) != text:
                errors.append(f"{reel.reel_id}: spoken body still requires punctuation mutation")
            spoken = validate_spoken_export_text(text)
            errors.extend(f"{reel.reel_id}: {error}" for error in spoken.errors)
            leaks = find_forbidden_substrings(text)
            errors.extend(f"{reel.reel_id}: forbidden metadata {leak}" for leak in leaks)

            nonempty = [line.strip() for line in text.splitlines() if line.strip()]
            if len(nonempty) >= 5:
                one_word = [line for line in nonempty if len(line.split()) == 1]
                if len(one_word) / len(nonempty) >= 0.4:
                    errors.append(f"{reel.reel_id}: word-per-line teleprompter layout")
            for line in nonempty:
                words = line.split()
                if len(words) == 1 and line.casefold() not in _ONE_WORD_EXCEPTIONS:
                    errors.append(f"{reel.reel_id}: unjustified one-word line '{line}'")
                if len(words) > 24:
                    errors.append(
                        f"{reel.reel_id}: dense spoken line has {len(words)} words"
                    )
        if module.module_project is not None:
            _validate_project_spoken_text(
                module.module_project.brief,
                f"module project {module.module_id}",
                errors,
            )

    if course.graduation_project is not None:
        _validate_project_spoken_text(
            course.graduation_project.brief,
            "graduation project",
            errors,
        )

    if errors:
        raise DocxVerificationError("; ".join(errors))


def _validate_project_spoken_text(
    text: str,
    label: str,
    errors: list[str],
) -> None:
    if not (text or "").strip():
        errors.append(f"{label}: empty spoken project instruction")
        return
    if strip_punctuation_from_spoken_body(text) != text:
        errors.append(f"{label}: instruction still requires punctuation mutation")
    leaks = find_forbidden_substrings(text)
    errors.extend(f"{label}: forbidden metadata {leak}" for leak in leaks)
    for line in [line.strip() for line in text.splitlines() if line.strip()]:
        if len(line.split()) > 24:
            errors.append(f"{label}: dense spoken line")


def verify_docx_archive(path: Path, course: FinalCourse) -> DocxVerificationReport:
    """Open ZIP/XML, headings, visible text, RTL/LTR, leaks, and project order."""
    report = DocxVerificationReport()
    if not path.exists() or not zipfile.is_zipfile(path):
        report.errors.append("not a valid DOCX ZIP archive")
        return report

    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            missing = sorted(_REQUIRED_DOCX_PARTS - names)
            if missing:
                report.errors.append("missing OOXML parts: " + ", ".join(missing))
            broken = archive.testzip()
            if broken:
                report.errors.append(f"corrupt ZIP member: {broken}")
            for name in names:
                if name.endswith(".xml"):
                    try:
                        ET.fromstring(archive.read(name))
                    except ET.ParseError as exc:
                        report.errors.append(f"invalid XML {name}: {exc}")
    except (OSError, zipfile.BadZipFile) as exc:
        report.errors.append(f"DOCX archive open failed: {exc}")
        return report

    document = Document(str(path))
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    visible = "\n".join(paragraph.text for paragraph in paragraphs)
    leaks = find_forbidden_substrings(visible)
    report.errors.extend(f"visible metadata leak: {leak}" for leak in leaks)
    if document.tables:
        report.errors.append("teleprompter DOCX contains a table")
    if any(header.text.strip() for section in document.sections for header in section.header.paragraphs):
        report.errors.append("teleprompter DOCX contains a running header")
    if any(footer.text.strip() for section in document.sections for footer in section.footer.paragraphs):
        report.errors.append("teleprompter DOCX contains a running footer")

    expected_headings: list[tuple[str, str]] = [("Title", course.title)]
    for module_index, module in enumerate(course.modules, start=1):
        expected_headings.append(
            ("Heading 1", module_heading(course, module_index, module.title))
        )
        for lesson_index, reel in enumerate(module.reels, start=1):
            expected_headings.append(
                ("Heading 2", lesson_heading(course, lesson_index, reel.title))
            )
        if module.module_project is not None:
            expected_headings.append(
                ("Heading 2", module.module_project.name or "مشروع الموديول")
            )
    if course.graduation_project is not None:
        expected_headings.append(
            ("Heading 1", course.graduation_project.name or "مشروع التخرج")
        )
    actual_headings = [
        (paragraph.style.name, paragraph.text)
        for paragraph in paragraphs
        if paragraph.style.name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
    ]
    if actual_headings != expected_headings:
        report.errors.append("heading hierarchy/order differs from approved course")

    expected_visible: list[str] = [course.title]
    for module_index, module in enumerate(course.modules, start=1):
        expected_visible.append(module_heading(course, module_index, module.title))
        for lesson_index, reel in enumerate(module.reels, start=1):
            expected_visible.append(lesson_heading(course, lesson_index, reel.title))
            expected_visible.extend(
                line.strip()
                for line in reel.script_text.splitlines()
                if line.strip()
            )
        if module.module_project is not None:
            expected_visible.append(module.module_project.name or "مشروع الموديول")
            if module.module_project.brief:
                expected_visible.append(module.module_project.brief)
    if course.graduation_project is not None:
        expected_visible.append(course.graduation_project.name or "مشروع التخرج")
        if course.graduation_project.brief:
            expected_visible.append(course.graduation_project.brief)
    actual_visible = [paragraph.text for paragraph in paragraphs]
    expected_text = "\n".join(expected_visible)
    actual_text = "\n".join(actual_visible)
    report.accepted_text_sha256 = _sha256(expected_text)
    report.docx_text_sha256 = _sha256(actual_text)
    if actual_visible != expected_visible:
        report.errors.append("visible DOCX paragraphs differ from accepted export content")
    if report.docx_text_sha256 != report.accepted_text_sha256:
        report.errors.append("rendered text fingerprint differs from accepted text")

    accepted_spoken = [
        _normalized(reel.script_text)
        for module in course.modules
        for reel in module.reels
    ]
    for reel_text in accepted_spoken:
        if reel_text and reel_text not in _normalized(visible):
            report.errors.append("accepted lesson text is missing, changed, or truncated")

    project_positions: list[int] = []
    module_positions: list[int] = []
    for index, paragraph in enumerate(paragraphs):
        if paragraph.style.name == "Heading 1":
            module_positions.append(index)
        if any(
            paragraph.text == (module.module_project.name or "مشروع الموديول")
            for module in course.modules
            if module.module_project is not None
        ):
            project_positions.append(index)
    for index, project_position in enumerate(project_positions):
        if index + 1 < len(module_positions) and project_position > module_positions[index + 1]:
            report.errors.append("module project is not between its module and the next module")

    if _course_is_arabic(course):
        for paragraph in paragraphs:
            if not _ARABIC.search(paragraph.text):
                continue
            p_pr = paragraph._p.pPr  # noqa: SLF001
            if p_pr is None or p_pr.find(qn("w:bidi")) is None:
                report.errors.append(f"Arabic paragraph lacks RTL bidi: {paragraph.text[:50]}")
                break

    return report


def render_docx_pages(
    path: Path,
    *,
    output_dir: Path,
    soffice_path: str | None = None,
) -> DocxVerificationReport:
    """Render every DOCX page and run automated clipping/overlap checks."""
    report = DocxVerificationReport()
    executable = soffice_path or shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        report.errors.append("DOCX visual renderer is unavailable")
        return report

    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="rukn-docx-render-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        completed = subprocess.run(
            [
                executable,
                f"-env:UserInstallation={profile_uri}",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=180,
            env={**os.environ, "HOME": profile},
        )
        pdf_path = output_dir / f"{path.stem}.pdf"
        if completed.returncode != 0 or not pdf_path.exists() or not pdf_path.stat().st_size:
            report.errors.append(
                "DOCX render failed: "
                + (completed.stderr or completed.stdout or "no PDF produced")[-500:]
            )
            return report

    with fitz.open(pdf_path) as pdf:
        report.page_count = pdf.page_count
        if not report.page_count:
            report.errors.append("rendered PDF has no pages")
        for page_index, page in enumerate(pdf):
            if not page.get_text("text").strip():
                report.errors.append(f"rendered page {page_index + 1} is blank")
            page_rect = page.rect
            blocks = [
                fitz.Rect(block[:4])
                for block in page.get_text("blocks")
                if str(block[4]).strip()
            ]
            for block in blocks:
                if not page_rect.contains(block):
                    report.errors.append(
                        f"rendered page {page_index + 1} has clipped text"
                    )
                    break
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            png_path = output_dir / f"page-{page_index + 1}.png"
            pixmap.save(png_path)
            report.rendered_page_paths.append(png_path)
    return report

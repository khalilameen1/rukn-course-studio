"""DOCX export for a finished, reviewed course (docs/ARCHITECTURE.md Stage 8).

Renders a `FinalCourse` (see app/schemas/generation.py) into a real .docx
following the teleprompter contract: course title, module headings, lesson
headings, spoken beats (no punctuation in body), module projects between
modules, and graduation project. Never production notes, sources, reviews,
Hook/Loop labels, or critic metadata.
"""

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from app.config import settings
from app.generation.contracts.spoken_final_master import (
    beats_to_plain_script,
    strip_punctuation_from_spoken_body,
    text_to_spoken_beats,
)
from app.schemas.generation import (
    CourseMap,
    FinalCourse,
    FinalModule,
    FinalReel,
    GeneratedReel,
    ModuleProject,
)
from app.services.docx_verification import (
    assert_final_course_ready_for_docx,
    lesson_heading,
    module_heading,
    render_docx_pages,
    verify_docx_archive,
)

TELEPROMPTER_FONT = "Arial"
BODY_FONT_SIZE = Pt(14)
BODY_LINE_SPACING = 1.5
_HEADING_STYLE_SIZES = {
    "Title": Pt(28),
    "Heading 1": Pt(20),
    "Heading 2": Pt(16),
    "Heading 3": Pt(14),
}
_HEADING_STYLE_SPACING = {
    "Title": (Pt(0), Pt(12)),
    "Heading 1": (Pt(18), Pt(10)),
    "Heading 2": (Pt(14), Pt(8)),
    "Heading 3": (Pt(10), Pt(5)),
}
_LATIN_TOKEN = re.compile(r"[A-Za-z0-9][A-Za-z0-9._+/#:@-]*")
_ARABIC = re.compile(r"[\u0600-\u06FF]")


def _set_style_font(style, size: Pt) -> None:
    style.font.name = TELEPROMPTER_FONT
    style.font.size = size
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:cs"), TELEPROMPTER_FONT)
    r_fonts.set(qn("w:ascii"), TELEPROMPTER_FONT)
    r_fonts.set(qn("w:hAnsi"), TELEPROMPTER_FONT)
    r_fonts.set(qn("w:eastAsia"), TELEPROMPTER_FONT)
    style.font.color.rgb = RGBColor(0, 0, 0)


def _set_run_direction(run, *, rtl: bool) -> None:
    r_pr = run._r.get_or_add_rPr()  # noqa: SLF001
    rtl_node = r_pr.find(qn("w:rtl"))
    if rtl_node is None:
        rtl_node = r_pr.makeelement(qn("w:rtl"), {})
        r_pr.append(rtl_node)
    rtl_node.set(qn("w:val"), "1" if rtl else "0")
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = r_pr.makeelement(qn("w:rFonts"), {})
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), TELEPROMPTER_FONT)


def _add_directional_runs(paragraph: Paragraph, text: str) -> None:
    cursor = 0
    for match in _LATIN_TOKEN.finditer(text or ""):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            _set_run_direction(run, rtl=bool(_ARABIC.search(run.text)))
        run = paragraph.add_run(match.group(0))
        _set_run_direction(run, rtl=False)
        cursor = match.end()
    if cursor < len(text or ""):
        run = paragraph.add_run(text[cursor:])
        _set_run_direction(run, rtl=bool(_ARABIC.search(run.text)))


def _set_rtl(paragraph: Paragraph) -> Paragraph:
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(p_pr.makeelement(qn("w:bidi"), {}))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        _set_run_direction(run, rtl=not bool(_LATIN_TOKEN.fullmatch(run.text.strip())))
    return paragraph


def _set_ltr(paragraph: Paragraph) -> Paragraph:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        _set_run_direction(run, rtl=False)
    return paragraph


def _add_text_paragraph(
    document: DocxDocument,
    text: str,
    *,
    force_rtl: bool,
    style: str | None = None,
) -> Paragraph:
    paragraph = document.add_paragraph(style=style)
    _add_directional_runs(paragraph, text)
    paragraph.paragraph_format.keep_together = True
    if force_rtl or _ARABIC.search(text or ""):
        return _set_rtl(paragraph)
    return _set_ltr(paragraph)


def _add_heading(
    document: DocxDocument,
    text: str,
    *,
    level: int,
    force_rtl: bool,
    page_break_before: bool = False,
) -> Paragraph:
    style = "Title" if level == 0 else f"Heading {level}"
    paragraph = _add_text_paragraph(
        document,
        text,
        force_rtl=force_rtl,
        style=style,
    )
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    return paragraph


def _apply_teleprompter_formatting(document: DocxDocument) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    _set_style_font(styles["Normal"], BODY_FONT_SIZE)
    styles["Normal"].paragraph_format.line_spacing = BODY_LINE_SPACING
    styles["Normal"].paragraph_format.space_after = Pt(8)
    for style_name, size in _HEADING_STYLE_SIZES.items():
        _set_style_font(styles[style_name], size)
        style_p_pr = styles[style_name].element.get_or_add_pPr()
        borders = style_p_pr.find(qn("w:pBdr"))
        if borders is not None:
            style_p_pr.remove(borders)
        before, after = _HEADING_STYLE_SPACING[style_name]
        styles[style_name].paragraph_format.space_before = before
        styles[style_name].paragraph_format.space_after = after
        styles[style_name].paragraph_format.keep_with_next = True

    # The teleprompter contract has no visible or embedded author metadata.
    core = document.core_properties
    core.author = ""
    core.last_modified_by = ""
    core.title = ""
    core.subject = ""
    core.comments = ""
    core.keywords = ""
    core.category = ""


def _spoken_body_lines(reel: FinalReel) -> list[str]:
    """Spoken lines for DOCX. Keeps intentional blank pause paragraphs."""
    beats = list(reel.spoken_beats or [])
    if beats:
        plain = beats_to_plain_script(beats)
    else:
        plain = reel.script_text or ""
    cleaned = strip_punctuation_from_spoken_body(plain)
    from app.generation.web_research import strip_research_leaks_from_script

    cleaned = strip_research_leaks_from_script(cleaned)
    # Preserve blank lines (teleprompter pauses); drop only trailing empties later.
    return [ln.strip() if ln.strip() else "" for ln in cleaned.splitlines()]


def _add_script_lines(document: DocxDocument, script_text: str) -> None:
    """Legacy helper — prefer reel-aware path; strips punctuation."""
    from app.generation.web_research import strip_research_leaks_from_script

    cleaned = strip_research_leaks_from_script(script_text or "")
    cleaned = strip_punctuation_from_spoken_body(cleaned)
    for line in cleaned.splitlines():
        if line.strip():
            _set_rtl(document.add_paragraph(line.strip()))
        else:
            _set_rtl(document.add_paragraph(""))


def _add_reel_script(
    document: DocxDocument,
    reel: FinalReel,
    *,
    force_rtl: bool,
) -> None:
    lines = _spoken_body_lines(reel)
    # Trim leading/trailing blank pause lines only.
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    if not lines:
        return
    for line in lines:
        _add_text_paragraph(document, line, force_rtl=force_rtl)


def _add_module_project(
    document: DocxDocument,
    project: ModuleProject,
    *,
    force_rtl: bool,
    heading_level: int = 2,
    page_break_before: bool = False,
) -> None:
    """Render a module/graduation project — not a numbered lesson."""
    _add_heading(
        document,
        project.name or "مشروع الموديول",
        level=heading_level,
        force_rtl=force_rtl,
        page_break_before=page_break_before,
    )
    if project.brief:
        _add_text_paragraph(document, project.brief, force_rtl=force_rtl)
    if project.closure:
        _add_text_paragraph(document, project.closure, force_rtl=force_rtl)
    # inputs/files, rubric, pass criteria, and skills tested are internal
    # project artefacts. Only the spoken project instruction belongs here.


def render_final_course_docx(final_course: FinalCourse) -> DocxDocument:
    """Build the in-memory python-docx Document."""
    document = Document()
    _apply_teleprompter_formatting(document)
    force_rtl = bool(
        final_course.thesis is not None
        and final_course.thesis.student_language.lower().startswith("ar")
    ) or bool(_ARABIC.search(final_course.title or ""))

    _add_heading(
        document,
        final_course.title or "Untitled Course",
        level=0,
        force_rtl=force_rtl,
    )

    for module_index, module in enumerate(final_course.modules, start=1):
        _add_heading(
            document,
            module_heading(final_course, module_index, module.title),
            level=1,
            force_rtl=force_rtl,
            page_break_before=module_index > 1,
        )

        for lesson_index, reel in enumerate(module.reels, start=1):
            _add_heading(
                document,
                lesson_heading(final_course, lesson_index, reel.title),
                level=2,
                force_rtl=force_rtl,
            )
            _add_reel_script(document, reel, force_rtl=force_rtl)

        # Only structured module_project is exported. Legacy bridge_project stays internal.
        if module.module_project is not None:
            _add_module_project(
                document,
                module.module_project,
                force_rtl=force_rtl,
            )

    return document


def extract_plain_text(document: DocxDocument) -> str:
    return "\n".join(p.text for p in document.paragraphs)


def next_version_number(existing_version_numbers: list[int]) -> int:
    return max(existing_version_numbers, default=0) + 1


def export_final_course_to_docx(
    final_course: FinalCourse, course_id: int, version_number: int
) -> Path:
    assert_final_course_ready_for_docx(final_course)
    document = render_final_course_docx(final_course)

    course_dir = settings.storage_outputs_dir / str(course_id)
    course_dir.mkdir(parents=True, exist_ok=True)
    output_path = course_dir / f"course_v{version_number}.docx"
    try:
        document.save(output_path)
        archive_report = verify_docx_archive(output_path, final_course)
        archive_report.raise_if_invalid()
        if settings.docx_visual_qa_required:
            with tempfile.TemporaryDirectory(
                prefix="rukn-docx-qa-", dir=course_dir
            ) as qa_dir:
                render_report = render_docx_pages(
                    output_path,
                    output_dir=Path(qa_dir),
                )
                render_report.raise_if_invalid()
    except Exception:
        output_path.unlink(missing_ok=True)
        raise
    return output_path


PARTIAL_DRAFT_NOTE = "Partial draft — generation stopped before completion."


def _assemble_partial_course(course_map: CourseMap, reels: list[GeneratedReel]) -> FinalCourse:
    sections: list[str] = []
    final_modules: list[FinalModule] = []

    for module in course_map.modules:
        module_reels = [r for r in reels if r.module_id == module.module_id]
        if not module_reels:
            continue

        sections.append(f"# {module.title}")
        final_reels: list[FinalReel] = []
        for reel in module_reels:
            sections.append(f"## {reel.title}")
            sections.append(reel.script_text)
            final_reels.append(
                FinalReel(
                    reel_id=reel.reel_id,
                    title=reel.title,
                    script_text=reel.script_text,
                    spoken_beats=list(reel.spoken_beats or []),
                    delivery_mode=reel.delivery_mode,
                    quality_status=reel.quality_status,
                )
            )

        final_modules.append(
            FinalModule(
                module_id=module.module_id,
                title=module.title,
                bridge_project=module.bridge_project,
                module_project=module.module_project,
                reels=final_reels,
            )
        )

    return FinalCourse(
        title=course_map.course_title,
        modules=final_modules,
        full_text="\n\n".join(sections),
        graduation_project=course_map.graduation_project,
        thesis=course_map.thesis,
    )


def build_partial_course_from_job(
    course_map_json: dict | str | None, completed_reels_json: list | str | None
) -> FinalCourse:
    from app.services.json_coerce import coerce_json_dict, coerce_json_list

    course_map = CourseMap.model_validate(coerce_json_dict(course_map_json) or {})
    reels = [
        GeneratedReel.model_validate(r)
        for r in coerce_json_list(completed_reels_json)
        if isinstance(r, dict)
    ]
    return _assemble_partial_course(course_map, reels)


def render_partial_course_docx(final_course: FinalCourse) -> DocxDocument:
    document = Document()
    _apply_teleprompter_formatting(document)

    force_rtl = bool(_ARABIC.search(final_course.title or ""))
    _add_text_paragraph(document, PARTIAL_DRAFT_NOTE, force_rtl=force_rtl)
    _add_heading(
        document,
        final_course.title or "Untitled Course",
        level=0,
        force_rtl=force_rtl,
    )

    for module_index, module in enumerate(final_course.modules, start=1):
        _add_heading(
            document,
            module_heading(final_course, module_index, module.title),
            level=1,
            force_rtl=force_rtl,
            page_break_before=module_index > 1,
        )

        for lesson_index, reel in enumerate(module.reels, start=1):
            _add_heading(
                document,
                lesson_heading(final_course, lesson_index, reel.title),
                level=2,
                force_rtl=force_rtl,
            )
            _add_reel_script(document, reel, force_rtl=force_rtl)

        # Only structured module_project is exported. Legacy bridge_project stays internal.
        if module.module_project is not None:
            _add_module_project(
                document,
                module.module_project,
                force_rtl=force_rtl,
            )

    return document


def export_partial_course_to_docx(final_course: FinalCourse, course_id: int, job_id: int) -> Path:
    document = render_partial_course_docx(final_course)

    course_dir = settings.storage_outputs_dir / str(course_id)
    course_dir.mkdir(parents=True, exist_ok=True)
    output_path = course_dir / f"partial_job_{job_id}.docx"
    document.save(output_path)
    return output_path

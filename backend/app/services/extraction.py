"""Source text extraction (docs/ARCHITECTURE.md §5 Source Ingestion Pipeline).

V1 supports .txt, .md, .docx, .pdf only. No OCR, no audio/video, no YouTube.

Every extraction attempt returns an `ExtractionResult` with one of the
statuses in app.services.source_status - callers must not treat
`extracted_text` as usable unless `status == READY` (or, if they choose to
allow it, `POOR_EXTRACTION` - never for password_required,
extraction_blocked, scanned_no_text, or failed, where `text` is always None).
"""

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.services.source_status import (
    EXTRACTION_BLOCKED,
    FAILED,
    PASSWORD_REQUIRED,
    POOR_EXTRACTION,
    READY,
    SCANNED_NO_TEXT,
)

# Below this many characters, text is treated as too thin to be useful.
MIN_ACCEPTABLE_CHARS = 40

# Crude "is this garbled" heuristic: real prose is mostly letters/spaces;
# a low alphabetic ratio usually means encoding noise or extraction garbage.
MIN_ALPHA_RATIO = 0.5


@dataclass
class ExtractionResult:
    status: str
    text: str | None = None
    error: str | None = None


def extract_text(path: Path, suffix: str, password: str | None = None) -> ExtractionResult:
    """Extract text from `path`. `suffix` must include the leading dot."""
    suffix = suffix.lower()
    try:
        if suffix in (".txt", ".md"):
            return _extract_plain_text(path)
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix == ".pdf":
            return _extract_pdf(path, password)
    except Exception as exc:  # noqa: BLE001 - any parser failure -> FAILED, never a crash
        return ExtractionResult(status=FAILED, error=str(exc))

    return ExtractionResult(status=FAILED, error=f"Unsupported extension '{suffix}'")


def _finalize(text: str) -> ExtractionResult:
    """Shared "is this text actually usable" check for every format."""
    cleaned = text.strip()
    if not cleaned:
        return ExtractionResult(status=EXTRACTION_BLOCKED)
    if _looks_garbled(cleaned):
        return ExtractionResult(status=POOR_EXTRACTION, text=cleaned)
    return ExtractionResult(status=READY, text=cleaned)


def _looks_garbled(text: str) -> bool:
    if len(text) < MIN_ACCEPTABLE_CHARS:
        return True
    alpha_count = sum(1 for c in text if c.isalpha())
    return (alpha_count / len(text)) < MIN_ALPHA_RATIO


def _extract_plain_text(path: Path) -> ExtractionResult:
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="utf-8", errors="replace")
    return _finalize(raw)


def _extract_docx(path: Path) -> ExtractionResult:
    document = DocxDocument(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return _finalize("\n".join(parts))


def _extract_pdf(path: Path, password: str | None) -> ExtractionResult:
    doc = fitz.open(str(path))
    try:
        if doc.needs_pass:
            if not password:
                return ExtractionResult(status=PASSWORD_REQUIRED)
            # Only ever tries the password the caller supplied - never
            # attempts to crack or bypass PDF protection.
            if not doc.authenticate(password):
                return ExtractionResult(
                    status=PASSWORD_REQUIRED, error="Incorrect password"
                )

        text_parts: list[str] = []
        has_any_images = False
        for page in doc:
            page_text = page.get_text().strip()
            if page_text:
                text_parts.append(page_text)
            elif page.get_images():
                has_any_images = True

        full_text = "\n\n".join(text_parts).strip()

        if not full_text:
            # Distinguish "looks scanned" from a generic block so the user
            # gets an accurate reason (see SOURCE_STATUS_MESSAGES).
            if has_any_images:
                return ExtractionResult(status=SCANNED_NO_TEXT)
            return ExtractionResult(status=EXTRACTION_BLOCKED)

        return _finalize(full_text)
    finally:
        doc.close()
